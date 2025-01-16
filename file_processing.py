import os
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from zipfile import BadZipFile

def process_pdf(file_path, chunk_size=500, chunk_overlap=100):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    documents = [Document(page_content=text)]
    return RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap).split_documents(documents)

def process_docx(file_path, chunk_size=500, chunk_overlap=100):
    try:
        doc = DocxDocument(file_path)  # Используем правильный класс
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)
        documents = [Document(page_content=text)]  # Создаем объект Document из langchain
        return RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap).split_documents(documents)
    except BadZipFile:
        raise ValueError("The DOCX file is corrupted or not a valid DOCX file.")
    except KeyError as e:
        raise ValueError(f"There is no item named '{e.args[0]}' in the archive")

def process_txt(file_path, chunk_size=500, chunk_overlap=100):
    with open(file_path, "r", encoding="utf-8") as file:
        text = file.read()
    documents = [Document(page_content=text)]
    return RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap).split_documents(documents)

def process_file(file_path, chunk_size=500, chunk_overlap=100):
    file_extension = os.path.splitext(file_path)[-1].lower()
    if file_extension == ".pdf":
        return process_pdf(file_path, chunk_size, chunk_overlap)
    elif file_extension == ".docx":
        return process_docx(file_path, chunk_size, chunk_overlap)
    elif file_extension == ".txt":
        return process_txt(file_path, chunk_size, chunk_overlap)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

if __name__ == "__main__":
    try:
        result = process_file("/Users/leonidstepanov/Desktop/site 2/Uploads/4.pdf", 500, 100)
        if result:
            print(result[0].page_content)
        else:
            print("No content processed.")
    except Exception as e:
        print(f"Ошибка при обработке файла DOCX: {e}")