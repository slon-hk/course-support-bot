import os
from typing import List, Dict, Union, Any
import PyPDF2
from docx import Document
import logging

logger = logging.getLogger(__name__)

def process_file(file_path: str) -> List[Dict[str, Any]]:
    """
    Обработать файл и извлечь из него текст.
    Поддерживает форматы PDF и DOCX.
    """
    try:
        logger.info(f"Начало обработки файла: {file_path}")
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == '.pdf':
            result = process_pdf(file_path)
            logger.info(f"PDF файл обработан, извлечено {len(result)} документов")
            return result
        elif file_extension in ['.docx', '.doc']:
            result = process_docx(file_path)
            logger.info(f"DOCX файл обработан, извлечено {len(result)} документов")
            return result
        else:
            logger.error(f"Неподдерживаемый формат файла: {file_extension}")
            return []

    except Exception as e:
        logger.error(f"Ошибка при обработке файла {file_path}: {str(e)}")
        return []

def process_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Извлечь текст из PDF файла"""
    try:
        documents = []
        with open(file_path, 'rb') as file:
            try:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                logger.info(f"Начало обработки PDF файла, всего страниц: {total_pages}")

                for page_num in range(total_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text and text.strip():
                            documents.append({
                                'text': text.strip(),
                                'page': page_num + 1,
                                'source': file_path
                            })
                            logger.debug(f"Успешно обработана страница {page_num + 1}")
                    except Exception as page_error:
                        logger.error(f"Ошибка при обработке страницы {page_num + 1}: {str(page_error)}")
                        continue

                return documents
            except Exception as pdf_error:
                logger.error(f"Ошибка при чтении PDF файла: {str(pdf_error)}")
                return []

    except Exception as e:
        logger.error(f"Ошибка при открытии PDF файла {file_path}: {str(e)}")
        return []

def process_docx(file_path: str) -> List[Dict[str, Any]]:
    """Извлечь текст из DOCX файла"""
    try:
        documents = []
        doc = Document(file_path)
        logger.info(f"Начало обработки DOCX файла, всего параграфов: {len(doc.paragraphs)}")

        for para_num, paragraph in enumerate(doc.paragraphs):
            try:
                text = paragraph.text.strip()
                if text:
                    documents.append({
                        'text': text,
                        'paragraph': para_num + 1,
                        'source': file_path
                    })
                    logger.debug(f"Успешно обработан параграф {para_num + 1}")
            except Exception as para_error:
                logger.error(f"Ошибка при обработке параграфа {para_num + 1}: {str(para_error)}")
                continue

        return documents
    except Exception as e:
        logger.error(f"Ошибка при обработке DOCX файла {file_path}: {str(e)}")
        return []